
from __future__ import annotations

import io
import json
import re
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any

import pandas as pd
import streamlit as st
from docx import Document


# ============================================================
# Настройки страницы
# ============================================================

st.set_page_config(
    page_title="Станок нарезки БЗ",
    page_icon="🧩",
    layout="wide",
)


# ============================================================
# Справочники UI → internal values
# ============================================================

CONTENT_TYPE_OPTIONS = {
    "Статья / инструкция": "article",
    "Сценарий / скрипт": "scenario",
    "Справочная информация": "reference_info",
    "Справочная карточка": "reference_article",
    "Быстрый ответ": "quick_answer",
    "Термин / глоссарий": "glossary_term",
    "Документ-оригинал": "source_document",
    "Архив": "archive",
    "На разбор": "review",
}

CONTENT_TYPE_LABELS = {internal: label for label, internal in CONTENT_TYPE_OPTIONS.items()}

RISK_OPTIONS = {
    "Низкий": "low",
    "Средний": "medium",
    "Высокий": "high",
}

PRIORITY_OPTIONS = ["P1", "P2", "P3", "P4", "P5"]

STATUS_OPTIONS = {
    "Черновик": "draft",
    "Нужно проверить": "needs_review",
    "Проверено": "approved",
    "Готово к импорту": "ready_for_import",
    "Опубликовано": "published",
    "Архив": "archived",
}

MANDATORY_OPTIONS = {
    "Нет": "no",
    "Да": "yes",
}


INTERNAL_COLUMNS = [
    "id",
    "block_id",
    "content_code",
    "content_type",
    "content_type_label",
    "title",
    "alternative_titles",
    "annotation",
    "section",
    "role",
    "process",
    "system",
    "risk",
    "priority",
    "tags",
    "owner",
    "reviewer",
    "mandatory_reading",
    "access_roles",
    "related_content_codes",
    "body_short",
    "body_for_whom",
    "body_when_apply",
    "body_steps",
    "body_important",
    "body_forbidden",
    "body_errors",
    "body_source",
    "status",
]


# ============================================================
# Helpers
# ============================================================

def normalize_spaces(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def safe_filename(text: str, max_len: int = 90) -> str:
    text = text.strip()
    text = re.sub(r'[\\/:*?"<>|]+', "_", text)
    text = re.sub(r"\s+", "_", text)
    text = re.sub(r"_+", "_", text).strip("_")
    return text[:max_len] or "untitled"


def get_label_by_value(options: dict[str, str], value: str, default_label: str | None = None) -> str:
    for label, internal in options.items():
        if internal == value:
            return label
    return default_label or next(iter(options.keys()))


def ensure_session() -> None:
    defaults = {
        "source_name": "",
        "source_text": "",
        "markers": [],
        "blocks": [],
        "cards": {},
        "selected_block_id": None,
        "project_loaded": False,
    }

    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


def read_docx_text(file_bytes: bytes) -> str:
    """
    Читает DOCX как текст.
    Вытаскивает обычные абзацы и простые таблицы.
    Таблицы превращает в строки с разделителем " | ".
    """
    document = Document(io.BytesIO(file_bytes))

    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = normalize_spaces(paragraph.text)
        if text:
            parts.append(text)

    # Таблицы добавляем в конец. Это не идеально, но лучше, чем потерять их совсем.
    # Для сложных таблиц всё равно нужен ручной разбор.
    for table_idx, table in enumerate(document.tables, start=1):
        parts.append(f"\n[ТАБЛИЦА {table_idx}]")
        for row in table.rows:
            cells = [normalize_spaces(cell.text) for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return normalize_spaces("\n".join(parts))


def parse_markers(text: str) -> list[str]:
    markers = []
    for line in text.splitlines():
        line = normalize_spaces(line)
        if not line:
            continue
        # Поддержка формата "заголовок, type=article" пока не режет по type,
        # но маркер берется до первой служебной запятой.
        marker = line.split(", type=")[0].strip()
        marker = marker.split(", тип=")[0].strip()
        if marker:
            markers.append(marker)
    return markers


def split_text_by_markers(source_text: str, markers: list[str]) -> tuple[list[dict[str, Any]], list[str]]:
    """
    Режем текст строго по ручным маркерам.
    Каждый маркер режется до следующего маркера.
    """
    errors: list[str] = []
    blocks: list[dict[str, Any]] = []

    if not source_text.strip():
        return [], ["Исходный текст пустой."]

    if not markers:
        return [], ["Список маркеров пустой."]

    lower_text = source_text.lower()
    positions: list[tuple[int, str]] = []

    for marker in markers:
        marker_low = marker.lower()
        pos = lower_text.find(marker_low)
        if pos == -1:
            errors.append(f"Маркер не найден: {marker}")
            continue

        # Проверка дублей: если marker встречается несколько раз, берем первое,
        # но подсвечиваем проблему.
        second_pos = lower_text.find(marker_low, pos + len(marker_low))
        if second_pos != -1:
            errors.append(f"Маркер найден больше одного раза, взято первое совпадение: {marker}")

        positions.append((pos, marker))

    positions = sorted(positions, key=lambda x: x[0])

    # Убираем дубликаты позиций
    unique_positions = []
    seen_pos = set()
    for pos, marker in positions:
        if pos in seen_pos:
            errors.append(f"Два маркера указывают на одну позицию: {marker}")
            continue
        seen_pos.add(pos)
        unique_positions.append((pos, marker))

    for idx, (start_pos, marker) in enumerate(unique_positions):
        end_pos = unique_positions[idx + 1][0] if idx + 1 < len(unique_positions) else len(source_text)
        raw = normalize_spaces(source_text[start_pos:end_pos])

        if len(raw) < 50:
            errors.append(f"Очень короткий блок после маркера: {marker}")

        block_id = f"BLK-{idx + 1:04d}"
        blocks.append(
            {
                "block_id": block_id,
                "order": idx + 1,
                "marker": marker,
                "raw_text": raw,
                "char_count": len(raw),
                "status": "new",
            }
        )

    return blocks, errors


def default_card_for_block(block: dict[str, Any], source_name: str) -> dict[str, Any]:
    raw_title = normalize_spaces(block.get("marker", ""))
    block_id = block["block_id"]

    return {
        "id": "",
        "block_id": block_id,
        "content_code": "",
        "content_type": "article",
        "content_type_label": "Статья / инструкция",
        "title": raw_title,
        "alternative_titles": "",
        "annotation": "",
        "section": "",
        "role": "",
        "process": "",
        "system": "",
        "risk": "medium",
        "priority": "P3",
        "tags": "",
        "owner": "",
        "reviewer": "",
        "mandatory_reading": "no",
        "access_roles": "",
        "related_content_codes": "",
        "body_short": "",
        "body_for_whom": "",
        "body_when_apply": "",
        "body_steps": block.get("raw_text", ""),
        "body_important": "",
        "body_forbidden": "",
        "body_errors": "",
        "body_source": f"{source_name} / {block_id} / marker: {raw_title}",
        "status": "draft",
    }


def build_content_code(card: dict[str, Any], index: int) -> str:
    if card.get("content_code"):
        return card["content_code"].strip()

    section = (card.get("section") or "").lower()
    ctype = card.get("content_type") or "article"

    if "мис" in section or "медиалог" in section:
        domain = "MIS"
    elif "дмс" in section or "страх" in section:
        domain = "DMS"
    elif "лаборатор" in section or "анализ" in section:
        domain = "LAB"
    elif "crm" in section or "usd" in section:
        domain = "CRM"
    elif "лояль" in section:
        domain = "LOY"
    elif "сервис" in section:
        domain = "SRV"
    elif "термин" in section or ctype == "glossary_term":
        domain = "TERM"
    else:
        domain = "ADM"

    type_map = {
        "article": "INS",
        "scenario": "SCN",
        "reference_info": "REF",
        "quick_answer": "QA",
        "glossary_term": "TERM",
        "reference_article": "REF",
        "source_document": "DOC",
        "archive": "ARC",
        "review": "REV",
    }
    return f"{domain}-{type_map.get(ctype, 'INS')}-{index:03d}"


def render_article_docx(card: dict[str, Any], raw_text: str | None = None) -> bytes:
    """
    Генерирует чистый DOCX для импорта в портал как статью.
    Тип объекта теперь явно попадает в документ.
    """
    doc = Document()

    code = card.get("content_code") or ""
    title = card.get("title") or "Без названия"
    content_type = card.get("content_type", "")
    content_type_label = card.get("content_type_label") or CONTENT_TYPE_LABELS.get(content_type, content_type)

    doc.add_heading(f"{code}. {title}".strip(". "), level=1)

    # Метаданные для редактора/импорта. Это видно в Word.
    doc.add_paragraph(f"Тип объекта: {content_type_label}")
    if card.get("section"):
        doc.add_paragraph(f"Раздел: {card.get('section')}")
    if card.get("role"):
        doc.add_paragraph(f"Роль: {card.get('role')}")
    if card.get("process"):
        doc.add_paragraph(f"Процесс: {card.get('process')}")
    if card.get("system"):
        doc.add_paragraph(f"Система: {card.get('system')}")
    if card.get("priority"):
        doc.add_paragraph(f"Приоритет: {card.get('priority')}")
    if card.get("tags"):
        doc.add_paragraph(f"Теги: {card.get('tags')}")

    sections = [
        ("Кратко", card.get("body_short")),
        ("Для кого", card.get("body_for_whom")),
        ("Когда применять", card.get("body_when_apply")),
        ("Порядок действий / содержание", card.get("body_steps")),
        ("Важно", card.get("body_important")),
        ("Что нельзя делать", card.get("body_forbidden")),
        ("Частые ошибки", card.get("body_errors")),
        ("Связанные материалы", card.get("related_content_codes")),
        ("Источник", card.get("body_source")),
    ]

    for heading, value in sections:
        value = normalize_spaces(str(value or ""))
        if not value:
            continue

        doc.add_heading(heading, level=2)

        # Если текст похож на список — сохраняем построчно.
        lines = [line.strip() for line in value.splitlines() if line.strip()]
        if len(lines) > 1:
            for line in lines:
                doc.add_paragraph(line)
        else:
            doc.add_paragraph(value)

    # Сырой текст как служебный блок можно оставить только для review.
    if card.get("status") == "needs_review" and raw_text:
        doc.add_heading("Сырой текст блока", level=2)
        doc.add_paragraph(raw_text)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_content_map(cards: list[dict[str, Any]]) -> bytes:
    rows = []
    for idx, card in enumerate(cards, start=1):
        card = dict(card)
        card["content_type_label"] = card.get("content_type_label") or CONTENT_TYPE_LABELS.get(card.get("content_type", ""), card.get("content_type", ""))
        row = {col: card.get(col, "") for col in INTERNAL_COLUMNS}
        if not row["id"]:
            row["id"] = f"KC-{idx:04d}"
        if not row["content_code"]:
            row["content_code"] = build_content_code(card, idx)
        rows.append(row)

    df = pd.DataFrame(rows, columns=INTERNAL_COLUMNS)

    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="content_map")

        ws = writer.book["content_map"]
        ws.freeze_panes = "A2"

        for cell in ws[1]:
            cell.font = cell.font.copy(bold=True)

        width_map = {
            "A": 12, "B": 14, "C": 18, "D": 18, "E": 35,
            "F": 40, "G": 45, "H": 24, "I": 20, "J": 24,
            "K": 24, "L": 12, "M": 10, "N": 30, "O": 24,
            "P": 24, "Q": 18, "R": 24, "S": 35, "T": 50,
            "U": 35, "V": 35, "W": 70, "X": 50, "Y": 50,
            "Z": 50, "AA": 45, "AB": 18,
        }

        for col, width in width_map.items():
            ws.column_dimensions[col].width = width

    return buffer.getvalue()


def build_export_zip() -> bytes:
    blocks = st.session_state.blocks
    cards_dict = st.session_state.cards

    cards: list[dict[str, Any]] = []
    block_by_id = {b["block_id"]: b for b in blocks}

    for idx, block in enumerate(blocks, start=1):
        block_id = block["block_id"]
        card = cards_dict.get(block_id) or default_card_for_block(block, st.session_state.source_name)
        card = dict(card)
        card["content_type_label"] = card.get("content_type_label") or CONTENT_TYPE_LABELS.get(card.get("content_type", ""), card.get("content_type", ""))
        if not card.get("id"):
            card["id"] = f"KC-{idx:04d}"
        if not card.get("content_code"):
            card["content_code"] = build_content_code(card, idx)
        cards.append(card)

    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as z:
        # content_map.xlsx
        z.writestr("metadata/content_map.xlsx", export_content_map(cards))

        # project.json
        project = {
            "source_name": st.session_state.source_name,
            "markers": st.session_state.markers,
            "blocks": st.session_state.blocks,
            "cards": cards,
            "exported_at": datetime.now().isoformat(timespec="seconds"),
        }
        z.writestr("metadata/project.json", json.dumps(project, ensure_ascii=False, indent=2))

        # cards as DOCX
        for idx, card in enumerate(cards, start=1):
            block = block_by_id.get(card["block_id"], {})
            raw_text = block.get("raw_text", "")
            code = card.get("content_code") or build_content_code(card, idx)
            title = safe_filename(card.get("title") or "Без названия")
            filename = f"{code}_{title}.docx"

            ctype = card.get("content_type", "article")
            if ctype == "scenario":
                folder = "scenarios_docx"
            elif ctype == "quick_answer":
                folder = "quick_answers_docx"
            elif ctype == "glossary_term":
                folder = "glossary_docx"
            elif ctype in {"reference_info", "reference_article"}:
                folder = "reference_docx"
            elif ctype in {"archive", "source_document"}:
                folder = "documents_or_archive_docx"
            elif ctype == "review":
                folder = "review_docx"
            else:
                folder = "articles_docx"

            z.writestr(f"{folder}/{filename}", render_article_docx(card, raw_text=raw_text))

        # raw blocks
        for block in blocks:
            filename = f"{block['block_id']}_{safe_filename(block.get('marker', 'block'))}.txt"
            z.writestr(f"raw_blocks/{filename}", block.get("raw_text", ""))

    return zip_buffer.getvalue()


def load_project_from_json(project_bytes: bytes) -> None:
    data = json.loads(project_bytes.decode("utf-8"))

    st.session_state.source_name = data.get("source_name", "")
    st.session_state.markers = data.get("markers", [])
    st.session_state.blocks = data.get("blocks", [])

    cards = {}
    for card in data.get("cards", []):
        block_id = card.get("block_id")
        if block_id:
            cards[block_id] = card

    st.session_state.cards = cards
    if st.session_state.blocks:
        st.session_state.selected_block_id = st.session_state.blocks[0]["block_id"]


# ============================================================
# UI
# ============================================================

ensure_session()

st.title("🧩 Станок нарезки базы знаний")
st.caption("Полуавтомат: человек задаёт смысл и поля, скрипт режет, оформляет и выгружает пакет для портала.")

tabs = st.tabs([
    "1. Исходник",
    "2. Маркеры",
    "3. Карточки",
    "4. Экспорт",
])


# ============================================================
# Tab 1 — Source
# ============================================================

with tabs[0]:
    st.header("1. Загрузка исходного документа")

    col1, col2 = st.columns([1, 1])

    with col1:
        uploaded = st.file_uploader(
            "Загрузить DOCX",
            type=["docx"],
            help="На первом этапе работаем с DOCX. PDF-сканы лучше заводить как документы-оригиналы.",
        )

        if uploaded is not None:
            file_bytes = uploaded.read()
            try:
                text = read_docx_text(file_bytes)
                st.session_state.source_name = uploaded.name
                st.session_state.source_text = text
                st.success(f"Файл загружен: {uploaded.name}")
            except Exception as exc:
                st.error(f"Не удалось прочитать DOCX: {exc}")

    with col2:
        project_upload = st.file_uploader(
            "Восстановить проект из project.json",
            type=["json"],
            help="project.json лежит внутри export ZIP в папке metadata.",
        )

        if project_upload is not None:
            try:
                load_project_from_json(project_upload.read())
                st.success("Проект восстановлен.")
            except Exception as exc:
                st.error(f"Не удалось восстановить проект: {exc}")

    if st.session_state.source_text:
        st.subheader("Предпросмотр текста")
        st.write(
            {
                "Файл": st.session_state.source_name,
                "Символов": len(st.session_state.source_text),
                "Строк": len(st.session_state.source_text.splitlines()),
            }
        )
        st.text_area(
            "Текст исходника",
            st.session_state.source_text[:50000],
            height=500,
            help="Показываются первые 50 000 символов, чтобы браузер не начал страдать.",
        )


# ============================================================
# Tab 2 — Markers
# ============================================================

with tabs[1]:
    st.header("2. Ручные маркеры разрезки")

    st.info(
        "Вставь заголовки блоков по одному на строку. "
        "Скрипт режет текст от текущего маркера до следующего. "
        "Это надежнее, чем пытаться угадать заголовки."
    )

    default_markers = "\n".join(st.session_state.markers)

    markers_text = st.text_area(
        "Маркеры / заголовки блоков",
        value=default_markers,
        height=350,
        placeholder=(
            "Алгоритм работы с пациентом за наличный расчет\n"
            "Алгоритм работы по сбору документов при первичном обращении\n"
            "ДОГОВОР-ОФЕРТЫ\n"
            "Алгоритм работы при обращении Клиента без предварительной записи"
        ),
    )

    col1, col2, col3 = st.columns([1, 1, 2])

    with col1:
        if st.button("Разрезать документ", type="primary"):
            markers = parse_markers(markers_text)
            st.session_state.markers = markers
            blocks, errors = split_text_by_markers(st.session_state.source_text, markers)
            st.session_state.blocks = blocks

            # Создаем дефолтные карточки только для новых блоков.
            for block in blocks:
                block_id = block["block_id"]
                if block_id not in st.session_state.cards:
                    st.session_state.cards[block_id] = default_card_for_block(
                        block,
                        st.session_state.source_name,
                    )

            if blocks:
                st.session_state.selected_block_id = blocks[0]["block_id"]

            if errors:
                st.warning("Разрезка выполнена, но есть предупреждения.")
                for err in errors:
                    st.write(f"⚠️ {err}")
            else:
                st.success(f"Готово. Блоков: {len(blocks)}")

    with col2:
        if st.button("Очистить блоки"):
            st.session_state.blocks = []
            st.session_state.cards = {}
            st.session_state.selected_block_id = None
            st.success("Блоки очищены.")

    if st.session_state.blocks:
        st.subheader("Результат нарезки")

        df_blocks = pd.DataFrame([
            {
                "block_id": b["block_id"],
                "order": b["order"],
                "marker": b["marker"],
                "char_count": b["char_count"],
            }
            for b in st.session_state.blocks
        ])

        st.dataframe(df_blocks, use_container_width=True, hide_index=True)


# ============================================================
# Tab 3 — Cards
# ============================================================

with tabs[2]:
    st.header("3. Карточки знания")

    if not st.session_state.blocks:
        st.warning("Сначала загрузите документ и разрежьте его по маркерам.")
    else:
        block_options = {
            f"{b['block_id']} — {b['marker'][:90]}": b["block_id"]
            for b in st.session_state.blocks
        }

        selected_label = st.selectbox(
            "Выбрать блок",
            options=list(block_options.keys()),
            index=0,
        )
        selected_block_id = block_options[selected_label]
        st.session_state.selected_block_id = selected_block_id

        block = next(b for b in st.session_state.blocks if b["block_id"] == selected_block_id)
        card = st.session_state.cards.get(selected_block_id) or default_card_for_block(
            block,
            st.session_state.source_name,
        )

        left, right = st.columns([1.05, 1.2])

        with left:
            st.subheader("Сырой блок")
            st.caption(f"{block['block_id']} · символов: {block['char_count']}")
            st.text_area(
                "Текст блока",
                value=block.get("raw_text", ""),
                height=760,
                key=f"raw_text_{selected_block_id}",
            )

        with right:
            st.subheader("Форма карточки")

            with st.form(key=f"card_form_{selected_block_id}"):
                st.markdown("### Основное")

                col_a, col_b = st.columns([1, 1])

                with col_a:
                    content_code = st.text_input(
                        "Артикул",
                        value=card.get("content_code", ""),
                        placeholder="ADM-INS-001",
                    )

                    content_type_label = st.selectbox(
                        "Тип объекта",
                        options=list(CONTENT_TYPE_OPTIONS.keys()),
                        index=list(CONTENT_TYPE_OPTIONS.values()).index(
                            card.get("content_type", "article")
                        )
                        if card.get("content_type", "article") in CONTENT_TYPE_OPTIONS.values()
                        else 0,
                    )

                    priority = st.selectbox(
                        "Приоритет",
                        options=PRIORITY_OPTIONS,
                        index=PRIORITY_OPTIONS.index(card.get("priority", "P3"))
                        if card.get("priority", "P3") in PRIORITY_OPTIONS
                        else 2,
                    )

                with col_b:
                    risk_label = st.selectbox(
                        "Риск",
                        options=list(RISK_OPTIONS.keys()),
                        index=list(RISK_OPTIONS.values()).index(card.get("risk", "medium"))
                        if card.get("risk", "medium") in RISK_OPTIONS.values()
                        else 1,
                    )

                    mandatory_label = st.selectbox(
                        "Обязательное ознакомление",
                        options=list(MANDATORY_OPTIONS.keys()),
                        index=list(MANDATORY_OPTIONS.values()).index(card.get("mandatory_reading", "no"))
                        if card.get("mandatory_reading", "no") in MANDATORY_OPTIONS.values()
                        else 0,
                    )

                    status_label = st.selectbox(
                        "Статус",
                        options=list(STATUS_OPTIONS.keys()),
                        index=list(STATUS_OPTIONS.values()).index(card.get("status", "draft"))
                        if card.get("status", "draft") in STATUS_OPTIONS.values()
                        else 0,
                    )

                title = st.text_input(
                    "Название",
                    value=card.get("title", ""),
                    placeholder="Первичный прием пациента за наличный расчет",
                )

                alternative_titles = st.text_area(
                    "Альтернативные заголовки / поисковые синонимы",
                    value=card.get("alternative_titles", ""),
                    height=80,
                    placeholder="первичный пациент; платный пациент; наличный расчет",
                )

                annotation = st.text_area(
                    "Аннотация",
                    value=card.get("annotation", ""),
                    height=100,
                    placeholder="Коротко: что это и когда использовать.",
                )

                st.markdown("### Классификация")

                col_c, col_d = st.columns([1, 1])

                with col_c:
                    section = st.text_input("Раздел", value=card.get("section", ""))
                    role = st.text_input("Роль", value=card.get("role", ""))
                    process = st.text_input("Процесс", value=card.get("process", ""))

                with col_d:
                    system = st.text_input("Система", value=card.get("system", ""))
                    tags = st.text_area(
                        "Теги",
                        value=card.get("tags", ""),
                        height=80,
                        placeholder="паспорт; согласие; касса; МИС",
                    )
                    access_roles = st.text_input(
                        "Роли доступа",
                        value=card.get("access_roles", ""),
                        placeholder="Администраторы; Руководители",
                    )

                st.markdown("### Ответственные и связи")

                col_e, col_f = st.columns([1, 1])

                with col_e:
                    owner = st.text_input("Владелец", value=card.get("owner", ""))
                    reviewer = st.text_input("Проверяющий", value=card.get("reviewer", ""))

                with col_f:
                    related_content_codes = st.text_area(
                        "Связанные материалы",
                        value=card.get("related_content_codes", ""),
                        height=100,
                        placeholder="MIS-INS-001; ADM-SCN-001; DOC-001",
                    )

                st.markdown("### Тело статьи / карточки")

                body_short = st.text_area("Кратко", value=card.get("body_short", ""), height=100)
                body_for_whom = st.text_area("Для кого", value=card.get("body_for_whom", ""), height=80)
                body_when_apply = st.text_area("Когда применять", value=card.get("body_when_apply", ""), height=90)
                body_steps = st.text_area(
                    "Порядок действий / содержание",
                    value=card.get("body_steps", ""),
                    height=260,
                )
                body_important = st.text_area("Важно", value=card.get("body_important", ""), height=100)
                body_forbidden = st.text_area("Что нельзя делать", value=card.get("body_forbidden", ""), height=100)
                body_errors = st.text_area("Частые ошибки", value=card.get("body_errors", ""), height=100)
                body_source = st.text_area("Источник", value=card.get("body_source", ""), height=80)

                submitted = st.form_submit_button("Сохранить карточку", type="primary")

                if submitted:
                    new_card = {
                        "id": card.get("id", ""),
                        "block_id": selected_block_id,
                        "content_code": content_code.strip(),
                        "content_type": CONTENT_TYPE_OPTIONS[content_type_label],
                        "content_type_label": content_type_label,
                        "title": title.strip(),
                        "alternative_titles": alternative_titles.strip(),
                        "annotation": annotation.strip(),
                        "section": section.strip(),
                        "role": role.strip(),
                        "process": process.strip(),
                        "system": system.strip(),
                        "risk": RISK_OPTIONS[risk_label],
                        "priority": priority,
                        "tags": tags.strip(),
                        "owner": owner.strip(),
                        "reviewer": reviewer.strip(),
                        "mandatory_reading": MANDATORY_OPTIONS[mandatory_label],
                        "access_roles": access_roles.strip(),
                        "related_content_codes": related_content_codes.strip(),
                        "body_short": body_short.strip(),
                        "body_for_whom": body_for_whom.strip(),
                        "body_when_apply": body_when_apply.strip(),
                        "body_steps": body_steps.strip(),
                        "body_important": body_important.strip(),
                        "body_forbidden": body_forbidden.strip(),
                        "body_errors": body_errors.strip(),
                        "body_source": body_source.strip(),
                        "status": STATUS_OPTIONS[status_label],
                    }

                    st.session_state.cards[selected_block_id] = new_card
                    st.success("Карточка сохранена.")

            st.divider()

            saved_card = st.session_state.cards.get(selected_block_id, card)
            docx_bytes = render_article_docx(saved_card, raw_text=block.get("raw_text", ""))
            code = saved_card.get("content_code") or "NO-CODE"
            filename = f"{code}_{safe_filename(saved_card.get('title') or 'card')}.docx"

            st.download_button(
                "Скачать DOCX этой карточки",
                data=docx_bytes,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )


# ============================================================
# Tab 4 — Export
# ============================================================

with tabs[3]:
    st.header("4. Экспорт пакета для портала")

    if not st.session_state.blocks:
        st.warning("Нет блоков для экспорта.")
    else:
        cards_list = []
        for idx, block in enumerate(st.session_state.blocks, start=1):
            card = st.session_state.cards.get(block["block_id"]) or default_card_for_block(
                block,
                st.session_state.source_name,
            )
            card = dict(card)
            if not card.get("id"):
                card["id"] = f"KC-{idx:04d}"
            if not card.get("content_code"):
                card["content_code"] = build_content_code(card, idx)
            cards_list.append(card)

        df_preview = pd.DataFrame(cards_list, columns=INTERNAL_COLUMNS)

        st.subheader("Предпросмотр content_map")
        st.dataframe(df_preview, use_container_width=True, hide_index=True)

        content_map_bytes = export_content_map(cards_list)

        col1, col2 = st.columns([1, 1])

        with col1:
            st.download_button(
                "Скачать content_map.xlsx",
                data=content_map_bytes,
                file_name="content_map.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

        with col2:
            zip_bytes = build_export_zip()
            st.download_button(
                "Скачать полный пакет ZIP",
                data=zip_bytes,
                file_name="portal_ready_package.zip",
                mime="application/zip",
                type="primary",
            )

        st.markdown("### Что внутри ZIP")

        st.code(
            """
metadata/content_map.xlsx
metadata/project.json
articles_docx/*.docx
scenarios_docx/*.docx
quick_answers_docx/*.docx
glossary_docx/*.docx
review_docx/*.docx
raw_blocks/*.txt
            """.strip()
        )
